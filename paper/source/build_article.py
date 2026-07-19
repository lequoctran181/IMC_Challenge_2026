#!/usr/bin/env python3
"""Build the Round 2 article from the organizer's DOCX template."""

from __future__ import annotations

import re
import sys
import argparse
from copy import deepcopy
from pathlib import Path

from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "source" / "Article template for IMC Challenge.docx"
MANUSCRIPT = ROOT / "manuscript.md"
OUTPUT = ROOT / "IMC_Challenge_Round2_NEU_AddictedTribes.docx"

NAVY = "17365D"
BLUE = "366092"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "666666"
BULLET_NUM_ID = None
EQUATION_NUMBER = 0


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row, repeat_header=False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    if repeat_header:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def create_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "360")
    indent.set(qn("w:hanging"), "180")
    p_pr.extend([tabs, indent])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    level.extend([start, num_fmt, level_text, level_jc, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("NEU.AddictedTribes  ·  Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def remove_template_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size=10.5, bold=False, italic=False, color="000000", font="Times New Roman") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


INLINE_RE = re.compile(
    r"(\\\(.+?\\\)|\[[^\]]+\]\(https?://[^)]+\)|`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)"
)


def add_hyperlink(paragraph, text: str, url: str, *, size=10.5, color=BLUE) -> None:
    """Add an external Word hyperlink without exposing a raw URL in the prose."""
    relation_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Times New Roman")
    run_properties.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    run_properties.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(round(size * 2)))
    run_properties.append(size_node)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def latex_to_omml(latex: str):
    """Convert LaTeX to a native, editable Word OMML equation object."""
    mathml = latex_to_mathml(latex)
    omml_text = mathml_to_omml(mathml)
    omml_text = omml_text.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1)
    return parse_xml(omml_text)


def add_inline(paragraph, text: str, size=10.5, italic=False, color="000000") -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, italic=italic, color=color)
        token = match.group(0)
        if token.startswith("\\("):
            paragraph._p.append(latex_to_omml(token[2:-2]))
        elif token.startswith("["):
            link = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2), size=size, color=BLUE)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(8.2, size - 1), font="Courier New", color=color)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, italic=italic, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, italic=italic, color=color)


def paragraph_base(paragraph, *, first_line=True, after=4, before=0, line=1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Inches(0.22)
    fmt.widow_control = True


def add_body(doc, text: str, *, italic=False, first_line=True, after=4) -> None:
    p = doc.add_paragraph()
    paragraph_base(p, first_line=first_line, after=after)
    add_inline(p, text, size=10.5, italic=italic)


def add_bullet(doc, text: str, *, compact=False) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(BULLET_NUM_ID))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1.4 if compact else 2.5)
    p.paragraph_format.line_spacing = 1.0 if compact else 1.04
    p.paragraph_format.widow_control = True
    add_inline(p, text, size=9.6 if compact else 10.25)


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.keep_with_next = True
    fmt.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "0" if level <= 2 else "1")
    p_pr.append(outline)
    if level == 1:
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(10)
        fmt.space_after = Pt(8)
        run = p.add_run(text)
        set_run_font(run, size=18, bold=True, color=NAVY)
    elif level == 2:
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, size=14, bold=True, color=NAVY)
    else:
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=BLUE)


def add_equation(doc, text: str) -> None:
    global EQUATION_NUMBER
    EQUATION_NUMBER += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(2.82), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.62), WD_TAB_ALIGNMENT.RIGHT)
    lead = p.add_run("\t")
    set_run_font(lead, size=10.5)
    p._p.append(latex_to_omml(text))
    tail = p.add_run(f"\t({EQUATION_NUMBER})")
    set_run_font(tail, size=9.5, color=MID_GRAY)


def add_reference(doc, text: str) -> None:
    """Typeset a bibliography entry with a conventional hanging indent."""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt.left_indent = Inches(0.25)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.02
    fmt.widow_control = True
    add_inline(p, text, size=9.7)


def add_code(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = p.add_run("\n".join(lines))
    set_run_font(run, size=8.1, font="Courier New")


def add_figure(doc, alt_caption: str, relative_path: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(ROOT / relative_path), width=Inches(5.65))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_together = True
    add_inline(cap, alt_caption, size=9.2, italic=True, color=MID_GRAY)


def add_table(doc, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    header = tuple(rows[0])
    prescribed = {
        ("Symbol", "Meaning"): [1.70, 3.95],
        ("Case", "Input V", "Final V'", "Retained", "Bottleneck and final strategy"): [0.95, 0.95, 0.80, 0.95, 2.00],
        ("Case", "Original V", "Final V'", "Retained", "Compression", "Score loss"): [1.00, 1.00, 0.80, 0.85, 1.00, 1.00],
        ("Submission / stage", "Official score", "Key change", "Output-count evidence"): [1.25, 1.00, 1.90, 1.50],
        ("Branch", "Base target / stage", "Normal term", "Additional term", "Offline tail"): [0.90, 1.20, 1.10, 1.20, 1.25],
        ("Case", "Input vertices", "Final vertices", "Retained", "Dominant bottleneck and final strategy"): [0.95, 0.90, 0.80, 0.85, 2.15],
        ("Case", "Original vertices", "Final vertices", "Retained", "Reduction", "Score loss"): [0.90, 0.85, 0.85, 1.15, 1.00, 0.90],
        ("Submission or stage", "Official score", "Main change", "Output-count evidence"): [1.15, 1.10, 1.95, 1.45],
        ("Branch", "Accepted frontier observation", "Rejected neighboring observation", "Interpretation"): [1.05, 1.35, 1.45, 1.80],
        ("Branch", "Base stage", "Normal term", "Additional mechanism", "Offline tail"): [0.95, 1.10, 1.20, 1.20, 1.20],
    }
    widths = prescribed.get(header)
    if widths is None:
        maxlens = [max(4, max(len(row[c]) if c < len(row) else 0 for row in rows)) for c in range(cols)]
        weights = [min(30, max(6, value)) for value in maxlens]
        total = sum(weights)
        widths = [5.65 * value / total for value in weights]
    descriptive_headers = {"Meaning", "Bottleneck and final strategy", "Key change", "Output-count evidence", "Base target / stage", "Normal term", "Additional term", "Offline tail"}
    # Word/LibreOffice honor the table grid before per-cell preferred widths.
    # Update both so wide narrative columns do not collapse into narrow strips.
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for c_idx, width in enumerate(widths):
        table.columns[c_idx].width = Inches(width)
        if c_idx < len(grid_cols):
            grid_cols[c_idx].set(qn("w:w"), str(int(Inches(width).twips)))
    for r_idx, (row_obj, values) in enumerate(zip(table.rows, rows)):
        prevent_row_split(row_obj, repeat_header=(r_idx == 0))
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c_idx, cell in enumerate(row_obj.cells):
            cell.width = Inches(widths[c_idx])
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is not None:
                tc_w.set(qn("w:w"), str(int(Inches(widths[c_idx]).twips)))
                tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if c_idx == 0 or header[c_idx] in descriptive_headers
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            text = values[c_idx] if c_idx < len(values) else ""
            table_font_size = 8.2 if cols >= 6 else 9.0
            add_inline(p, text, size=table_font_size, color="FFFFFF" if r_idx == 0 else "000000")
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build() -> None:
    global BULLET_NUM_ID, EQUATION_NUMBER
    EQUATION_NUMBER = 0
    doc = Document(str(TEMPLATE))
    remove_template_body(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.591)
    section.footer_distance = Inches(0.689)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    BULLET_NUM_ID = create_bullet_numbering(doc)
    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    add_page_number(footer.add_paragraph())

    core = doc.core_properties
    core.title = "Certified Perceptual Mesh Simplification under a 21-Second and 128-KiB Budget"
    core.subject = "IMC Challenge Round 2 Article — Problem B"
    core.author = "NEU.AddictedTribes"
    core.last_modified_by = "NEU.AddictedTribes"
    core.keywords = "mesh simplification, QEM, SSIM, renderer-aware optimization, IMC Challenge"
    core.comments = "Authoritative Round 2 article for Kattis submission 20082703 (93.830074)."

    lines = MANUSCRIPT.read_text(encoding="utf-8").splitlines()
    i = 0
    cover = True
    repeated_title_seen = False
    compact_bullets = False
    top_sections = {
        "Abstract", "1. Introduction", "2. Related Literature", "3. Methodology",
        "4. Results and Discussions", "5. Conclusion", "References",
    }
    page_break_sections = {"3. Methodology", "4. Results and Discussions"}

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "<!-- PAGE BREAK -->":
            doc.add_page_break()
            cover = False
            i += 1
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(72 if cover else 8)
            p.paragraph_format.space_after = Pt(16 if cover else 6)
            run = p.add_run(title)
            set_run_font(run, size=28 if cover else 23, bold=True, color=NAVY)
            if not cover:
                repeated_title_seen = True
            i += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading in page_break_sections:
                doc.add_page_break()
            add_heading(doc, heading, 1 if heading in top_sections else 2)
            i += 1
            continue
        if line.startswith("### "):
            heading = line[4:].strip()
            compact_bullets = heading.startswith("Appendix C.")
            add_heading(doc, heading, 2)
            i += 1
            continue
        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
            i += 1
            continue
        if line.startswith("!["):
            match = re.fullmatch(r"!\[(.+)\]\((.+)\)", line)
            if match:
                add_figure(doc, match.group(1), match.group(2))
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code(doc, code_lines)
            i += 1
            continue
        if line.startswith("$$") and line.endswith("$$"):
            add_equation(doc, line[2:-2].strip())
            i += 1
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip(), compact=compact_bullets)
            i += 1
            continue
        if re.match(r"^\[\d+\]\s", line):
            add_reference(doc, line)
            i += 1
            continue
        if cover and line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(36)
            run = p.add_run(line[1:-1])
            set_run_font(run, size=14, italic=True, color=MID_GRAY)
            i += 1
            continue
        if not cover and repeated_title_seen and line.startswith("*") and line.endswith("*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[1:-1])
            set_run_font(run, size=11.5, italic=True, color=MID_GRAY)
            repeated_title_seen = False
            i += 1
            continue
        if cover and line.startswith("Problem name:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(62)
            p.paragraph_format.space_after = Pt(10)
            label, value = line.split(":", 1)
            r1 = p.add_run(label + ":")
            set_run_font(r1, size=15, bold=True, color=NAVY)
            r2 = p.add_run(value)
            set_run_font(r2, size=15, italic=True)
            i += 1
            continue
        if cover and line.startswith("Team name:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label, value = line.split(":", 1)
            r1 = p.add_run(label + ":")
            set_run_font(r1, size=18, bold=True, color=NAVY)
            r2 = p.add_run(value)
            set_run_font(r2, size=18, italic=True)
            i += 1
            continue
        if cover and line.startswith("GitHub repository:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(4)
            label, value = line.split(":", 1)
            r1 = p.add_run(label + ": ")
            set_run_font(r1, size=11, bold=True, color=NAVY)
            match = re.search(r"\[([^\]]+)\]\((https?://[^)]+)\)", value)
            if match:
                add_hyperlink(p, match.group(1), match.group(2), size=11, color=BLUE)
            i += 1
            continue

        italic_body = False
        if line.startswith("*") and line.endswith("*"):
            italic_body = True
            line = line[1:-1]
        add_body(doc, line, italic=italic_body, first_line=not line.startswith("Authoritative artifact:"))
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


def main() -> None:
    global TEMPLATE, MANUSCRIPT, OUTPUT
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=TEMPLATE,
                        help="path to the organizer's DOCX template")
    parser.add_argument("--manuscript", type=Path, default=MANUSCRIPT,
                        help="Markdown source manuscript")
    parser.add_argument("--output", type=Path, default=OUTPUT,
                        help="output DOCX path")
    args = parser.parse_args()
    if not args.template.is_file():
        parser.error(f"template not found: {args.template}")
    if not args.manuscript.is_file():
        parser.error(f"manuscript not found: {args.manuscript}")

    TEMPLATE = args.template.resolve()
    MANUSCRIPT = args.manuscript.resolve()
    OUTPUT = args.output.resolve()
    build()


if __name__ == "__main__":
    main()
